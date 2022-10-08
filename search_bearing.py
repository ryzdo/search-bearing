from docx import Document
import re


def searcValue (text, regex):
    mo = regex.search(text)

    if mo:
        return mo.group().strip()


def isDate (text):
    '''Поиск даты в документе'''
    dateRegex = re.compile(r'\s\d\d\.\d\d\.\d{4}')
    return searcValue(text, dateRegex)
    
def isExecutor (text):
    '''Поиск Исполнителя в документе'''
    dateRegex = re.compile(r'.+сполнитель')
    return searcValue(text, dateRegex)

def isCustomer (text):
    '''Поиск Заказчика в документе'''
    dateRegex = re.compile(r'.+аказчик')
    return searcValue(text, dateRegex)

def isRecipient (text):
    '''Поиск Получателя в документе'''
    dateRegex = re.compile(r'.+олучатель')
    return searcValue(text, dateRegex)
# открываем существующий документ 
doc = Document('/home/denis/Документы/python_project/search-bearing/akt.docx')

act = None # акт
act_date = None # дата
executor = None # исполнитель
customer = None # заказчик
recipient = None # получатель
brand_size = None # бренд типоразмер 

bearing_number = None # номер подшипника
date_manufacture = None # дата изготовления
reason_transfer = None # причина передачи

print()
# Поиск в шапке документа
for paragraph in doc.paragraphs:

    # Поиск даты: если дата None - Искать
    if not act_date:
        act_date = isDate(paragraph.text)

    # поиск исполнителя
    if not executor:
        executor = isExecutor(paragraph.text)
    
    # поиск Зпкпзчика
    if not customer:
        customer = isCustomer(paragraph.text)

    # поиск получателя
    if not recipient:
        recipient = isRecipient(paragraph.text)


print(act_date)
print(executor)
print(customer)
print(recipient)
print(brand_size)

# table = doc.tables[1]

# # читаем данные из таблиц
# for table in doc.tables:
#     print(table)
#     for row in table.rows:
#         string = ''
#         for cell in row.cells:
#             string = string + cell.text + ' '
#         print(string)